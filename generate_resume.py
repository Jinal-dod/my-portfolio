import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_font(run, font_name='Times New Roman', size=11, bold=False, color=None):
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)

def add_heading(doc, text, size=12):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_font(run, size=size, bold=True, color=(44, 106, 171)) # Professional Blue
    p.space_after = Pt(2)
    p.space_before = Pt(10)

def add_highlighted_text(paragraph, text, highlights):
    """Adds text to a paragraph with specific words bolded."""
    words = text.split()
    i = 0
    while i < len(words):
        found = False
        for highlight in sorted(highlights, key=len, reverse=True):
            highlight_words = highlight.split()
            clean_text = " ".join(words[i:i+len(highlight_words)]).strip(",.")
            clean_highlight = " ".join(highlight_words).strip(",.")
            
            if clean_text == clean_highlight:
                run = paragraph.add_run(" " + " ".join(words[i:i+len(highlight_words)]) if paragraph.runs else " ".join(words[i:i+len(highlight_words)]))
                set_font(run, size=9, bold=True)
                i += len(highlight_words)
                found = True
                break
        if not found:
            run = paragraph.add_run(" " + words[i] if paragraph.runs else words[i])
            set_font(run, size=9, bold=False)
            i += 1

def create_resume():
    doc = Document()
    
    # Set margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)

    # Header
    p1 = doc.add_paragraph()
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run1 = p1.add_run("JINAL DOD")
    set_font(run1, size=24, bold=True, color=(44, 106, 171))
    p1.space_after = Pt(0)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run("Senior Odoo Developer | Python & JavaScript Expert | ERP Implementation Specialist | Integration Specialist")
    set_font(run2, size=10, bold=True)
    p2.space_after = Pt(0)

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run3 = p3.add_run("Ahmedabad, India 380058 | +91 8320708960 | jinaldod8151@gmail.com | www.linkedin.com/in/jinal-dod-8956b716a/")
    set_font(run3, size=9)
    p3.space_after = Pt(10)

    # SUMMARY
    add_heading(doc, "SUMMARY")
    p = doc.add_paragraph("Dynamic Senior Odoo Developer with 4.5+ years of expertise in architecting high-impact ERP solutions and custom module development. Proven track record in orchestrating complex API integrations and performance optimizations at scale, while leading high-performance teams to deliver mission-critical projects. Adept in advanced PostgreSQL scripting and passionate about crafting intuitive, user-centric experiences through cutting-edge UI/UX design.")
    set_font(p.runs[0], size=10)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # SKILLS
    add_heading(doc, "SKILLS")
    skills_categories = [
        ("ERP Development", "Odoo (v10–v19), Odoo.sh, Full ERP Cycle, Custom Modules, Odoo Frontend Development"),
        ("Databases", "PostgreSQL (Scripting), MySQL, Database Management"),
        ("Tools", "Git, VSCode, Antigravity, Cursor"),
        ("Programming", "Python, JavaScript, HTML5, CSS3"),
        ("Web Frameworks", "Django"),
        ("Others", "Workflow Automation, Bug Fixing, Performance Optimization, UI/UX Enhancement, API Integration, XML/JSON Handling, Technical Documentation, Code Optimization Strategies, Project Delivery and Mentoring")
    ]
    
    for cat, val in skills_categories:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.2)
        run_cat = p.add_run(cat + ": ")
        set_font(run_cat, size=9, bold=True)
        run_val = p.add_run(val)
        set_font(run_val, size=9)
        p.space_after = Pt(1)

    # EXPERIENCE
    add_heading(doc, "EXPERIENCE")
    
    highlights = [
        "design, development, and customization of Odoo ERP solutions",
        "Activity Calendar functionality",
        "eCommerce integration with Shopify and ShipStation",
        "3PL logistics workflows",
        "mobile-integrated damage return workflows",
        "BOM versioning and multi-level BOM consumption logic",
        "Python and PostgreSQL scripts for data migration and database correction",
        "document scanning module",
        "OCR/AI APIs",
        "sentiment analysis scores",
        "custom Odoo modules",
        "custom APIs and third-party integrations",
        "QWeb-based UI templates",
        "dedicated developer",
        "PSMR cycle automation",
        "advanced Odoo modules",
        "POS promotions, quick payment discounts, forecast-based reporting, and custom product packaging workflows",
        "Many2One (M2O) widget behavior",
        "custom automated actions, server actions, and scheduled jobs",
        "Klarna and Vipps",
        "Mentored junior developers, conducted trainings and technical interviews,",
        "Collaborated with functional teams",
        "Trained and mentored new joiners",
        "Conducted on-site client visits (Ahmedabad, Dubai)"
    ]

    experience_data = [
        {
            "dates": "09/2025 – Current",
            "title": "Senior Odoo Developer",
            "company": "Wanbuffer Services – Ahmedabad",
            "bullets": [
                "Led design, development, and customization of Odoo ERP solutions across Inventory, Manufacturing, Sales, Purchase, and Accounting modules.",
                "Developed advanced Activity Calendar functionality for improved task planning, scheduling, and operational tracking.",
                "Implemented end-to-end eCommerce integration with Shopify and ShipStation, including checkout process, order lifecycle, and automated customer synchronization via APIs and webhooks, shipment process.",
                "Designed and implemented complete 3PL logistics workflows covering inbound, storage, picking, packing, and outbound operations.",
                "Implemented mobile-integrated damage return workflows for real-time warehouse and QC operations.",
                "Developed BOM versioning and multi-level BOM consumption logic along with projection planning reports for manufacturing and demand forecasting.",
                "Created Python and PostgreSQL scripts for data migration and database correction in production environments.",
                "Mentored junior developers, conducted trainings and technical interviews, and supported team members through code reviews and debugging.",
                "Collaborated with functional teams for requirement gathering and ERP implementation ensuring scalable and business-aligned solutions.",
                "Designed and developed a document scanning module in Odoo to process Vendor Bills, Delivery Challans, and Resumes using OCR/AI APIs.",
                "Integrated external APIs to extract structured data and sentiment analysis scores from uploaded documents."
            ]
        },
        {
            "dates": "08/2023 – 08/2025",
            "title": "Odoo Developer",
            "company": "Kanak Infosystems LLP – Gandhinagar",
            "bullets": [
                "Designed and implemented custom Odoo modules across sales, purchase, accounting, inventory, and manufacturing, websites, POS.",
                "Enhanced ERP workflows through custom APIs and third-party integrations (payment gateways, e-commerce).",
                "Developed QWeb-based UI templates for intuitive and responsive user experiences.",
                "Led as a dedicated developer on key client projects, ensuring timely delivery and high-quality solutions.",
                "Trained and mentored new joiners on Odoo development best practices, coding standards, and workflows.",
                "Conducted on-site client visits (Ahmedabad, Dubai) for process mapping, requirement gathering, and PSMR cycle automation."
            ]
        },
        {
            "dates": "08/2021 – 08/2023",
            "title": "Odoo Developer",
            "company": "Odoo India Pvt. Ltd. – Gandhinagar",
            "bullets": [
                "Customized and developed advanced Odoo modules across versions 13 to 16, covering key areas such as Sales, Purchase, Accounting, Repair, Manufacturing, Inventory, and HR.",
                "Implemented value-adding business features such as POS promotions, quick payment discounts, forecast-based reporting, and custom product packaging workflows to improve operational efficiency.",
                "Enhanced Many2One (M2O) widget behavior with dynamic filters, contextual domain logic, and custom relational rules, improving user experience and data accuracy.",
                "Optimized workflows through implementation of custom automated actions, server actions, and scheduled jobs, minimizing manual intervention.",
                "Upgraded, integrated, and maintained third-party payment gateways, including Klarna and Vipps, ensuring seamless payment processing and compliance with updated API specifications."
            ]
        }
    ]

    for exp in experience_data:
        table = doc.add_table(rows=1, cols=2)
        table.autofit = False
        table.columns[0].width = Inches(1.5)
        table.columns[1].width = Inches(6.0)
        
        row = table.rows[0].cells
        p_date = row[0].paragraphs[0]
        run_date = p_date.add_run(exp['dates'])
        set_font(run_date, size=8, bold=True)
        
        cell_details = row[1]
        p_title = cell_details.paragraphs[0]
        run_title = p_title.add_run(exp['title'])
        set_font(run_title, size=10, bold=True)
        
        p_company = cell_details.add_paragraph()
        run_company = p_company.add_run(exp['company'])
        set_font(run_company, size=10, bold=True)

        for bullet in exp['bullets']:
            p_bullet = cell_details.add_paragraph(style='List Bullet')
            add_highlighted_text(p_bullet, bullet, highlights)
            p_bullet.paragraph_format.space_after = Pt(2)

    # EDUCATION
    add_heading(doc, "EDUCATION")
    edu_data = [
        {
            "dates": "2016 – 2020",
            "title": "Bachelor of Engineering (BE), Information Technology,",
            "school": "Shantilal Shah Government Engineering College, Bhavnagar (GTU)",
            "note": "CGPA: 7.30",
            "location": "Bhavnagar"
        },
        {
            "dates": "2014 – 2016",
            "title": "Higher Secondary Education,",
            "school": "Smt. R.C. Patel Secondary & Higher Secondary School, Gandhinagar",
            "note": "Percentage: 58.80%",
            "location": "Gandhinagar"
        }
    ]

    for edu in edu_data:
        table = doc.add_table(rows=1, cols=2)
        table.autofit = False
        table.columns[0].width = Inches(1.5)
        table.columns[1].width = Inches(6.0)
        row = table.rows[0].cells
        p_date = row[0].paragraphs[0]
        run_date = p_date.add_run(edu['dates'])
        set_font(run_date, size=8.5, bold=True)
        
        cell_details = row[1]
        p_title = cell_details.paragraphs[0]
        run_title = p_title.add_run(edu['title'])
        set_font(run_title, size=10.5, bold=True)
        
        p_school = cell_details.add_paragraph()
        run_school = p_school.add_run(edu['school'])
        set_font(run_school, size=10, bold=True)
        
        p_loc = cell_details.add_paragraph()
        run_loc = p_loc.add_run(edu['location'])
        set_font(run_loc, size=9)
        
        p_note = cell_details.add_paragraph()
        run_note = p_note.add_run(edu['note'])
        set_font(run_note, size=9.5)

    # WEBSITES
    add_heading(doc, "WEBSITES, PORTFOLIOS, PROFILES")
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run("linkedin.com/in/jinal-dod-8956b716a")
    set_font(run, size=9)

    # CERTIFICATIONS
    add_heading(doc, "CERTIFICATIONS")
    p = doc.add_paragraph("Python Programming & Django Framework, TOPS Technologies, Ahmedabad, Gujarat")
    set_font(p.runs[0], size=10)

    # LANGUAGES
    add_heading(doc, "LANGUAGES")
    for lang in ["Gujarati", "Hindi", "English"]:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(lang)
        set_font(run, size=9)

    doc.save("Jinal_Dod_Resume.docx")
    print("Resume generated: Jinal_Dod_Resume.docx")

if __name__ == "__main__":
    create_resume()
